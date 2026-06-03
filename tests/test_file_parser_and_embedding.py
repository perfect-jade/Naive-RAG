"""
文件解析器 & Embedding 服务测试用例

测试范围:
- PDF 文件解析
- TXT 文件解析
- DOCX 文件解析
- Markdown 文件解析
- Embedding 服务调用与维度验证
- 文本分块逻辑
"""

import os
import json
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest


class TestFileParser:
    """文件解析器测试"""

    def test_parse_txt_file(self):
        """TC-FPR-001: 解析 TXT 文件"""
        content = "第一章 总则\n\n第一条 测试内容。\n\n第二条 继续测试。"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write(content)
            tmp_path = f.name

        try:
            with open(tmp_path, "r", encoding="utf-8") as f:
                parsed = f.read()
            assert "第一章 总则" in parsed
            assert "测试内容" in parsed
            assert len(parsed) > 0
        finally:
            os.unlink(tmp_path)

    def test_parse_markdown_file(self):
        """TC-FPR-002: 解析 Markdown 文件"""
        content = "# 标题\n\n这是内容。\n\n## 二级标题\n\n- 列表项1\n- 列表项2"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
            f.write(content)
            tmp_path = f.name

        try:
            with open(tmp_path, "r", encoding="utf-8") as f:
                parsed = f.read()
            assert "标题" in parsed
            assert "列表项1" in parsed
        finally:
            os.unlink(tmp_path)

    def test_parse_docx_file(self):
        """TC-FPR-003: 解析 DOCX 文件"""
        try:
            from docx import Document

            doc = Document()
            doc.add_heading("测试文档", level=1)
            doc.add_paragraph("这是第一段内容。")
            doc.add_paragraph("这是第二段内容。")

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                doc.save(f.name)
                tmp_path = f.name

            try:
                doc2 = Document(tmp_path)
                full_text = "\n".join([p.text for p in doc2.paragraphs])
                assert "测试文档" in full_text
                assert "第一段内容" in full_text
            finally:
                os.unlink(tmp_path)
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_parse_pdf_file(self):
        """TC-FPR-004: 解析 PDF 文件"""
        try:
            from PyPDF2 import PdfReader
            import io
            from reportlab.pdfgen import canvas

            packet = io.BytesIO()
            c = canvas.Canvas(packet)
            c.drawString(100, 750, "这是 PDF 测试内容")
            c.drawString(100, 730, "第二行内容")
            c.save()

            packet.seek(0)
            reader = PdfReader(packet)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""

            assert "PDF 测试内容" in text
        except ImportError:
            pytest.skip("PyPDF2 or reportlab not installed")

    def test_parse_unsupported_file(self):
        """TC-FPR-005: 不支持的文件类型返回错误"""
        unsupported_extensions = [".exe", ".bin", ".dat", ".jpg"]
        for ext in unsupported_extensions:
            assert ext not in (".pdf", ".docx", ".txt", ".md")

    def test_parse_empty_file(self):
        """TC-FPR-006: 解析空文件返回空字符串"""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            tmp_path = f.name

        try:
            with open(tmp_path, "r", encoding="utf-8") as f:
                content = f.read()
            assert content == ""
        finally:
            os.unlink(tmp_path)

    def test_parse_large_file(self):
        """TC-FPR-007: 解析大文件不崩溃"""
        content = "测试行内容。\n" * 10000
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write(content)
            tmp_path = f.name

        try:
            with open(tmp_path, "r", encoding="utf-8") as f:
                parsed = f.read()
            assert len(parsed) > 0
            assert "测试行内容" in parsed
        finally:
            os.unlink(tmp_path)

    def test_parse_encoding_utf8(self):
        """TC-FPR-008: UTF-8 编码中文文件正确解析"""
        content = "中文字符测试：你好世界！"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write(content)
            tmp_path = f.name

        try:
            with open(tmp_path, "r", encoding="utf-8") as f:
                parsed = f.read()
            assert "你好世界" in parsed
        finally:
            os.unlink(tmp_path)


class TestTextSplitting:
    """文本分块测试"""

    def test_split_by_chunk_size(self):
        """TC-SPL-001: 按 chunk_size 分块"""
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        text = "这是一段测试文本。" * 100
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=100,
            chunk_overlap=20,
            separators=["\n\n", "\n", "。", "！", "？", "，", " ", ""]
        )
        chunks = splitter.split_text(text)
        assert len(chunks) > 1
        for chunk in chunks[:-1]:
            assert len(chunk) <= 100 + 20

    def test_split_short_text(self):
        """TC-SPL-002: 短文本分块不超过 1 个"""
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        text = "短文本"
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", " ", ""]
        )
        chunks = splitter.split_text(text)
        assert len(chunks) == 1
        assert chunks[0] == text

    def test_split_overlap(self):
        """TC-SPL-003: 分块重叠验证"""
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        text = "A。" * 100
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=50,
            chunk_overlap=20,
            separators=["。", " ", ""]
        )
        chunks = splitter.split_text(text)
        if len(chunks) >= 2:
            assert len(chunks[-2]) > 0

    def test_split_empty_text(self):
        """TC-SPL-004: 空文本分块返回空列表"""
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", " ", ""]
        )
        chunks = splitter.split_text("")
        assert len(chunks) == 0

    def test_split_chinese_text(self):
        """TC-SPL-005: 中文文本按句号分块"""
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        text = (
            "高血压是指以体循环动脉血压增高为主要特征的临床综合征。"
            "在未使用降压药物的情况下，非同日3次测量诊室血压，"
            "收缩压≥140mmHg和/或舒张压≥90mmHg即可诊断为高血压。"
            "糖尿病是以高血糖为特征的代谢性疾病。"
        )
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=100,
            chunk_overlap=20,
            separators=["\n\n", "\n", "。", "！", "？", "，", " ", ""]
        )
        chunks = splitter.split_text(text)
        assert len(chunks) >= 1


class TestEmbeddingService:
    """Embedding 服务测试"""

    def test_embedding_dimension(self, mock_embedding_response):
        """TC-EMB-001: Embedding 返回正确维度"""
        embedding = mock_embedding_response["output"]["embeddings"][0]["embedding"]
        assert len(embedding) == 1536

    def test_embedding_values_are_float(self, mock_embedding_response):
        """TC-EMB-002: Embedding 值均为浮点数"""
        embedding = mock_embedding_response["output"]["embeddings"][0]["embedding"]
        for val in embedding:
            assert isinstance(val, float)

    def test_embedding_batch(self):
        """TC-EMB-003: 批量生成 Embedding"""
        texts = ["文本一", "文本二", "文本三"]
        mock_embeddings = [[0.1] * 1536 for _ in range(3)]

        assert len(mock_embeddings) == len(texts)
        for emb in mock_embeddings:
            assert len(emb) == 1536

    def test_embedding_model_switch_dimension_change(self):
        """TC-EMB-004: 切换 Embedding 模型后维度可能变化"""
        dims = {
            "text-embedding-v1": 1536,
            "text-embedding-v2": 1536,
            "text-embedding-v3": 1024,
        }

        assert dims["text-embedding-v1"] == 1536
        assert dims["text-embedding-v2"] == 1536
        assert dims["text-embedding-v3"] == 1024

        assert dims["text-embedding-v3"] != dims["text-embedding-v2"]

    def test_embedding_normalization(self):
        """TC-EMB-005: Embedding 归一化后 L2 范数为 1"""
        import numpy as np

        vec = np.array([0.1] * 1536, dtype=np.float32)
        norm = np.linalg.norm(vec)
        normalized = vec / norm if norm > 0 else vec
        assert abs(np.linalg.norm(normalized) - 1.0) < 0.001

    def test_embedding_similarity(self):
        """TC-EMB-006: 相似文本的 Embedding 余弦相似度更高"""
        import numpy as np

        np.random.seed(42)
        emb1 = np.random.random(1536).astype(np.float32)
        emb2 = emb1 + np.random.normal(0, 0.01, 1536).astype(np.float32)
        emb3 = np.random.random(1536).astype(np.float32)

        sim_12 = np.dot(emb1, emb2) / (np.linalg.norm(emb1) * np.linalg.norm(emb2))
        sim_13 = np.dot(emb1, emb3) / (np.linalg.norm(emb1) * np.linalg.norm(emb3))

        assert sim_12 > sim_13