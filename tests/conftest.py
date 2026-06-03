import os
import sys
import json
import shutil
import tempfile
import uuid
from pathlib import Path
from unittest.mock import MagicMock, patch, AsyncMock

import pytest
import numpy as np
from fastapi.testclient import TestClient

sys.path.insert(0, str(Path(__file__).parent.parent / "backend"))

TEST_DATA_DIR = Path(__file__).parent / "test_data"


@pytest.fixture(scope="session")
def project_root():
    return Path(__file__).parent.parent


@pytest.fixture
def temp_data_dir():
    tmp = tempfile.mkdtemp(prefix="rag_test_")
    yield Path(tmp)
    shutil.rmtree(tmp, ignore_errors=True)


@pytest.fixture
def mock_embedding_response():
    return {
        "output": {
            "embeddings": [
                {"embedding": [0.1] * 1536, "text_index": 0}
            ]
        },
        "usage": {"total_tokens": 10},
        "request_id": str(uuid.uuid4()),
        "status_code": 200
    }


@pytest.fixture
def mock_llm_route_response_medical():
    return {
        "output": {
            "choices": [
                {
                    "message": {
                        "content": '{"industry": "医疗健康", "confidence": "high"}',
                        "role": "assistant"
                    },
                    "finish_reason": "stop"
                }
            ]
        },
        "usage": {"total_tokens": 50},
        "request_id": str(uuid.uuid4()),
        "status_code": 200
    }


@pytest.fixture
def mock_llm_route_response_general():
    return {
        "output": {
            "choices": [
                {
                    "message": {
                        "content": '{"industry": "general", "confidence": "low"}',
                        "role": "assistant"
                    },
                    "finish_reason": "stop"
                }
            ]
        },
        "usage": {"total_tokens": 50},
        "request_id": str(uuid.uuid4()),
        "status_code": 200
    }


@pytest.fixture
def mock_llm_chat_stream():
    def _generate():
        tokens = ["根据", "资料", "显示", "，", "高血压", "的", "诊断", "标准", "为", "收缩压≥140mmHg"]
        for token in tokens:
            yield {
                "output": {
                    "choices": [
                        {
                            "message": {"content": token, "role": "assistant"},
                            "finish_reason": "null"
                        }
                    ]
                },
                "usage": {"total_tokens": 1},
                "request_id": str(uuid.uuid4()),
                "status_code": 200
            }
        yield {
            "output": {
                "choices": [
                    {
                        "message": {"content": "", "role": "assistant"},
                        "finish_reason": "stop"
                    }
                ]
            },
            "usage": {"total_tokens": 1},
            "request_id": str(uuid.uuid4()),
            "status_code": 200
        }
    return _generate


@pytest.fixture
def mock_dashscope(mock_embedding_response, mock_llm_route_response_medical, mock_llm_chat_stream):
    with patch("dashscope.TextEmbedding") as mock_embed, \
         patch("dashscope.Generation") as mock_gen:

        mock_embed.call = MagicMock(return_value=MagicMock(
            output=MagicMock(embeddings=[
                MagicMock(embedding=[0.1] * 1536)
            ]),
            status_code=200
        ))

        mock_gen.call = MagicMock()

        yield {
            "embedding": mock_embed,
            "generation": mock_gen
        }


@pytest.fixture
def sample_industry_data():
    return {
        "name": "医疗健康",
        "description": "医疗健康领域专业知识库"
    }


@pytest.fixture
def sample_industry_data_2():
    return {
        "name": "法律法规",
        "description": "法律法规领域专业知识库"
    }


@pytest.fixture
def sample_knowledge_text():
    return {
        "title": "高血压诊疗指南",
        "content": "高血压是指以体循环动脉血压（收缩压和/或舒张压）增高为主要特征，可伴有心、脑、肾等器官的功能或器质性损害的临床综合征。在未使用降压药物的情况下，非同日3次测量诊室血压，收缩压≥140mmHg和/或舒张压≥90mmHg即可诊断为高血压。",
        "tags": ["心血管", "慢性病"]
    }


@pytest.fixture
def sample_knowledge_text_2():
    return {
        "title": "糖尿病防治手册",
        "content": "糖尿病是一种以高血糖为特征的代谢性疾病。高血糖则是由于胰岛素分泌缺陷或其生物作用受损，或两者兼有引起。长期存在的高血糖，导致各种组织，特别是眼、肾、心脏、血管、神经的慢性损害、功能障碍。",
        "tags": ["内分泌", "慢性病"]
    }


@pytest.fixture
def sample_chat_query():
    return {
        "query": "高血压的诊断标准是什么？",
        "history": [],
        "industry": None
    }


@pytest.fixture
def sample_chat_query_with_history():
    return {
        "query": "那糖尿病呢？",
        "history": [
            {"role": "user", "content": "高血压的诊断标准是什么？"},
            {"role": "assistant", "content": "高血压的诊断标准是收缩压≥140mmHg和/或舒张压≥90mmHg。"}
        ],
        "industry": "medical_health"
    }


@pytest.fixture
def sample_settings():
    return {
        "llm_model": "qwen-plus",
        "embedding_model": "text-embedding-v2",
        "api_key": None
    }


@pytest.fixture
def sample_settings_update():
    return {
        "llm_model": "qwen-max",
        "embedding_model": "text-embedding-v3",
        "api_key": "sk-test-key-12345"
    }


@pytest.fixture
def sample_pdf_path():
    path = TEST_DATA_DIR / "sample.pdf"
    TEST_DATA_DIR.mkdir(parents=True, exist_ok=True)
    if not path.exists():
        path.write_text("这是一个测试 PDF 文件的内容。", encoding="utf-8")
    return path


@pytest.fixture
def sample_txt_path():
    path = TEST_DATA_DIR / "sample.txt"
    TEST_DATA_DIR.mkdir(parents=True, exist_ok=True)
    if not path.exists():
        path.write_text(
            "第一章 总则\n\n第一条 为了保护公民、法人的合法权益，制定本法。\n\n第二条 本法适用于中华人民共和国境内。",
            encoding="utf-8"
        )
    return path


@pytest.fixture
def sample_docx_path():
    path = TEST_DATA_DIR / "sample.docx"
    TEST_DATA_DIR.mkdir(parents=True, exist_ok=True)
    if not path.exists():
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("测试文档", level=1)
            doc.add_paragraph("这是一个测试 Word 文档的内容。")
            doc.save(str(path))
        except ImportError:
            path.write_text("测试文档\n这是一个测试 Word 文档的内容。", encoding="utf-8")
    return path


@pytest.fixture
def faiss_dimension():
    return 1536


@pytest.fixture
def sample_embeddings(faiss_dimension):
    np.random.seed(42)
    return np.random.random((50, faiss_dimension)).astype(np.float32)


@pytest.fixture
def sample_chunks():
    return [
        {"id": str(uuid.uuid4()), "content": "高血压是指以体循环动脉血压增高为主要特征的临床综合征。"},
        {"id": str(uuid.uuid4()), "content": "收缩压≥140mmHg和/或舒张压≥90mmHg即可诊断为高血压。"},
        {"id": str(uuid.uuid4()), "content": "糖尿病是以高血糖为特征的代谢性疾病。"},
        {"id": str(uuid.uuid4()), "content": "胰岛素分泌缺陷或其生物作用受损导致高血糖。"},
        {"id": str(uuid.uuid4()), "content": "长期高血糖导致各种组织的慢性损害和功能障碍。"},
        {"id": str(uuid.uuid4()), "content": "为了保护公民法人的合法权益制定本法。"},
        {"id": str(uuid.uuid4()), "content": "本法适用于中华人民共和国境内。"},
        {"id": str(uuid.uuid4()), "content": "金融市场的稳定性对经济发展至关重要。"},
        {"id": str(uuid.uuid4()), "content": "股票投资需要关注市盈率和市净率等指标。"},
        {"id": str(uuid.uuid4()), "content": "风湿性关节炎是一种自身免疫性疾病。"},
    ]